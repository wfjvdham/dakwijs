# -*- coding: utf-8 -*-

# Run this app with `python app.py` and
# visit http://127.0.0.1:5050/ in your web browser.

import base64
import json
import math
import os
import sys
import subprocess

import dash
import dash_bootstrap_components as dbc
import dash_core_components as dcc
import dash_html_components as html
import dash_table
import flask
import pandas as pd
from PIL import Image
from dash.dependencies import Input, Output, State
from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches

df = pd.read_excel("./Solor 2021.xlsm", sheet_name=1, names=['id', 'desc', 'price'], usecols=[0, 1, 2],
                   dtype={'id': str, 'desc': str, 'price': str})
df['price'] = df['price'].str.replace('€', '').str.strip().str.replace(',', '.')
df = df.astype({'price': float})
df['count'] = 0

template_filename = "Solar template 2021.docx"
paneel_filename = 'paneel.png'
temp_advice_filename = 'temp_advies.docx'
image_filename = "image_advies.jpg"
advice_filename_docx = "downloads/advies.docx"
advice_filename_pdf = "downloads/advies.pdf"

app = dash.Dash(__name__, external_stylesheets=[dbc.themes.JOURNAL])

dropdown_card = dbc.Card(
    dbc.CardBody([
        dbc.FormGroup([
            dbc.Label('Indak / Opdak', html_for='daksysteem'),
            dcc.Dropdown(
                id='daksysteem',
                options=[
                    {'label': 'Indak', 'value': 'Indak'},
                    {'label': 'Opdak', 'value': 'Opdak'},
                ],
                value='Indak', clearable=False
            ),
        ]),
        dbc.FormGroup([
            dbc.Label('Landscape / Portrait', html_for='indeling'),
            dcc.Dropdown(
                id='indeling',
                options=[
                    {'label': 'Landscape', 'value': 'LND'},
                    {'label': 'Portrait', 'value': 'POR'},
                ],
                value='LND', clearable=False
            ),
        ]),
    ])
)

paneel_card = dbc.Card(
    dbc.CardBody([
        dbc.FormGroup([
            dbc.Label('Paneellengte [mm]', html_for='paneellengte'),
            dbc.Input(id='paneellengte', value=0, min=0, type='number'),
        ]),
        dbc.FormGroup([
            dbc.Label('Paneelbreedte [mm]', html_for='paneelbreedte'),
            dbc.Input(id='paneelbreedte', value=0, min=0, type='number'),
        ]),
        dbc.FormGroup([
            dbc.Label('Paneeldikte [mm]', html_for='paneeldikte'),
            dbc.Input(id='paneeldikte', value=0, min=0, type='number'),
        ]),
    ])
)

rijen_kolommen_card = dbc.Card(
    dbc.CardBody([
        dbc.FormGroup([
            dbc.Label('Aantal rijen', html_for='rijen'),
            dbc.Input(id='rijen', value=0, min=0, type='number'),
        ]),
        dbc.FormGroup([
            dbc.Label("Aantal kolommen", html_for='kolommen'),
            dbc.Input(id='kolommen', value=0, min=0, type='number'),
        ]),
    ])
)

rest_card = dbc.Card(
    dbc.CardBody([
        dbc.FormGroup([
            dbc.Label("Kleur frame", html_for='kleurFrame'),
            dcc.Dropdown(
                id='kleurFrame',
                options=[
                    {'label': 'ALU', 'value': 'ALU'},
                    {'label': 'ALU Zwart', 'value': 'ALZ'}
                ],
                value='ALU', clearable=False
            ),
        ]),
        dbc.FormGroup([
            dbc.Label("Plaat / Paneel", html_for='toepassing'),
            dcc.Dropdown(
                id='toepassing',
                options=[
                    {'label': 'Plaat', 'value': 'PLA'},
                    {'label': 'Paneel', 'value': 'PAN'}
                ],
                value='PAN', clearable=False
            ),
        ]),
        dbc.FormGroup([
            dbc.Label('Dakhelling'),
            dcc.Slider(
                id='dakhelling',
                min=0,
                max=90,
                value=0,
            )
        ])
    ])
)

input_tab = dbc.Card(
    dbc.CardBody(
        dbc.Row([
            dbc.Col(dropdown_card, sm=3),
            dbc.Col(paneel_card, sm=3),
            dbc.Col(rijen_kolommen_card, sm=3),
            dbc.Col(rest_card, sm=3)
        ])
    )
)

table_header = [
    html.Thead(html.Tr([html.Th("Naam"), html.Th("Waarde")]))
]
table_body = [html.Tbody([
    html.Tr([html.Td("Raillengte"), html.Td(3000, id='raillengte')]),
    html.Tr([html.Td("Rol"), html.Td("1140 x 10000")]),
    html.Tr([html.Td("Eindklem"), html.Td(40, id="eindklem")]),
    html.Tr([html.Td("Tussenklem"), html.Td(22, id="tussenklem")]),
    html.Tr([html.Td("Anker plaatsen om de"), html.Td(800, id="anker_plaatsen_om_de")]),
    html.Tr([html.Td("Benodigde overlap"), html.Td(50, id="benodigde_overlap")])
])]

constants_tab = dbc.Card(
    dbc.CardBody([
        dbc.Table(table_header + table_body, bordered=True, hover=True, responsive=True,
                  striped=True)
    ])
)

table_header = [
    html.Thead(html.Tr([html.Th("Naam"), html.Th("Waarde")]))
]
table_body = [html.Tbody([
    html.Tr([html.Td("Lengthe Rail"), html.Td(id='lengte_rail')]),
    html.Tr([html.Td("Aantal rijen rails"), html.Td(id="aantal_rijen_rails")]),
    html.Tr([html.Td("Totale lengte rails"), html.Td(id="totale_lengte_rails")]),
    html.Tr([html.Td("Aantal rails van 3 meter per rij"), html.Td(id="aantal_rails_van_3_meter_per_rij")]),
    html.Tr([html.Td("Lengte 1 rol"), html.Td(id="lengte_1_rol")]),
    html.Tr([html.Td("Breedte pv"), html.Td(id="breedte_pv")]),
    html.Tr([html.Td("Aantal rijen rollen"), html.Td(id="aantal_rijen_rollen")]),
    html.Tr([html.Td("Dakgoten"), html.Td(id="dakgoten")]),
    html.Tr([html.Td("Schuimstrook driehoek profiel"), html.Td(id="schuimstrook_driehoek_profiel")]),
    html.Tr([html.Td("Railverbinder"), html.Td(id="railverbinder")]),
    html.Tr([html.Td("Aantal ankers op 1 rail"), html.Td(id="aantal_ankers_op_1_rail")]),
    html.Tr([html.Td("Ankers"), html.Td(id="ankers")]),
    html.Tr([html.Td("Schroeven voor ankers"), html.Td(id="schroeven_voor_ankers")]),
    html.Tr([html.Td("Beugels"), html.Td(id="beugels")]),
    html.Tr([html.Td("Schroeven voor beugels"), html.Td(id="schroeven_voor_beugels")]),
    html.Tr([html.Td("Eindklemmen"), html.Td(id="eindklemmen")]),
    html.Tr([html.Td("Middenklemmen"), html.Td(id="middenklemmen")]),
    html.Tr([html.Td("Haak"), html.Td(id="haak")]),
    html.Tr([html.Td("Schroeven voor hoek"), html.Td(id="schroeven_voor_hoek")]),
    html.Tr([html.Td("Totaal aantal rails van 3m"), html.Td(id="totaal_aantal_rails_van_3m")]),
])]

results_tab = dbc.Card(
    dbc.CardBody([
        dbc.Table(table_header + table_body, bordered=True, hover=True, responsive=True,
                  striped=True),
        html.Div(id='data', style={'display': 'none'})
    ])
)

download_tab = dbc.Card(
    dbc.CardBody([
        dbc.FormGroup([
            dbc.Label("Referentie nr.", html_for='referentie_nr'),
            dbc.Input(id='referentie_nr', type='text'),
        ]),
        dbc.FormGroup([
            dbc.Label("Relatie", html_for='relatie'),
            dbc.Input(id='relatie', type='text'),
        ]),
        dbc.FormGroup([
            dbc.Label("Contactpersoon", html_for='contactpersoon'),
            dbc.Input(id='contactpersoon', type='text'),
        ]),
        dbc.FormGroup([
            dbc.Label("Project", html_for='project'),
            dbc.Input(id='project', type='text'),
        ]),
        dbc.FormGroup([
            dbc.Label("Partijen", html_for='partijen'),
            dbc.Input(id='partijen', type='text'),
        ]),
        dbc.FormGroup([
            dbc.Label("Adviseur", html_for='adviseur'),
            dbc.Input(id='adviseur', type='text'),
        ]),
        html.Button(
            'Maak Advies', id='create_advice', className='btn btn-secondary', n_clicks=0
        ),
        html.Br(),
        html.Hr(),
        html.A(
            id='download-link-docx', children='Download Advies (docx)',
            className='btn btn-primary', href='/{}'.format(advice_filename_docx), style={'display': 'none'}
        ),
        html.A(
            id='download-link-pdf', children='Download Advies (pdf)',
            className='btn btn-primary', href='/{}'.format(advice_filename_pdf), style={'display': 'none'}
        )
    ])
)

app.layout = dbc.Tabs(
    [
        dbc.Tab(input_tab, label="Invoer"),
        dbc.Tab(constants_tab, label="Constanten"),
        dbc.Tab(results_tab, label="Resultaten"),
        dbc.Tab([
            dbc.Col(
                html.Div(id='table', className="pt-3"),
                width={'size': 10, 'offset': 1}
            ),
            dbc.Col(
                html.Div(id='total_price', className="pt-3"),
                width={'size': 2, 'offset': 9}
            )
        ], label="Leverlijst"),
        dbc.Tab(
            html.Div(
                html.Div(
                    id='square', className="pt-5"
                ),
                className="row d-flex justify-content-center"
            ),
            label="Visual"
        ),
        dbc.Tab(download_tab, label="Download Advies")
    ]
)


@app.server.route('/downloads/<path:path>')
def serve_static(path):
    root_dir = os.getcwd()
    return flask.send_from_directory(
        os.path.join(root_dir, 'downloads'), path
    )


@app.callback(
    Output(component_id='lengte_rail', component_property='children'),
    Input(component_id='indeling', component_property='value'),
    Input(component_id='paneelbreedte', component_property='value'),
    Input(component_id='rijen', component_property='value'),
    Input(component_id='kolommen', component_property='value'),
    Input(component_id='tussenklem', component_property='children'),
    Input(component_id='eindklem', component_property='children'),
)
def update_output_div(indeling, paneelbreedte, rijen, kolommen, tussenklem, eindklem):
    if indeling == 'LND':
        result = rijen * paneelbreedte + ((rijen-1) * tussenklem) + 2 * eindklem + 50
    else:
        result = kolommen * paneelbreedte + ((kolommen - 1) * tussenklem) + 2 * eindklem + 50
    return result


@app.callback(
    Output(component_id='aantal_rijen_rails', component_property='children'),
    Input(component_id='indeling', component_property='value'),
    Input(component_id='rijen', component_property='value'),
    Input(component_id='kolommen', component_property='value'),
)
def update_output_div(indeling, rijen, kolommen):
    if indeling == 'LND':
        result = 2 * kolommen
    else:
        result = 2 * rijen
    return result


@app.callback(
    Output(component_id='totale_lengte_rails', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
    Input(component_id='lengte_rail', component_property='children'),
)
def update_output_div(aantal_rijen_rails, lengte_rail):
    result = aantal_rijen_rails * lengte_rail
    return result


@app.callback(
    Output(component_id='aantal_rails_van_3_meter_per_rij', component_property='children'),
    Input(component_id='lengte_rail', component_property='children'),
    Input(component_id='raillengte', component_property='children'),
)
def update_output_div(lengte_rail, raillengte):
    result = math.ceil(lengte_rail / raillengte)
    return result


@app.callback(
    Output(component_id='lengte_1_rol', component_property='children'),
    Input(component_id='kolommen', component_property='value'),
    Input(component_id='paneellengte', component_property='value'),
    Input(component_id='tussenklem', component_property='children'),
    Input(component_id='eindklem', component_property='children'),
    Input(component_id='lengte_rail', component_property='children'),
    Input(component_id='indeling', component_property='value')
)
def update_output_div(kolommen, paneellengte, tussenklem, eindklem, lengte_rail, indeling):
    if indeling == 'LND':
        result = ((kolommen * paneellengte + ((kolommen - 1) * tussenklem) + 2 * eindklem) - 100)
    else:
        result = lengte_rail - 100
    return result


@app.callback(
    Output(component_id='breedte_pv', component_property='children'),
    Input(component_id='rijen', component_property='value'),
    Input(component_id='paneelbreedte', component_property='value'),
    Input(component_id='paneellengte', component_property='value'),
    Input(component_id='tussenklem', component_property='children'),
    Input(component_id='eindklem', component_property='children'),
    Input(component_id='indeling', component_property='value')
)
def update_output_div(rijen, paneelbreedte, paneellengte, tussenklem, eindklem, indeling):
    if indeling == 'LND':
        result = (rijen * paneelbreedte + ((rijen - 1) * tussenklem) + 2 * eindklem)
    else:
        result = (rijen * paneellengte + ((rijen - 1) * eindklem) + 2 * eindklem)
    return result


@app.callback(
    Output(component_id='aantal_rijen_rollen', component_property='children'),
    Input(component_id='breedte_pv', component_property='children'),
    Input(component_id='tussenklem', component_property='children'),
)
def update_output_div(breedte_pv, tussenklem):
    result = math.ceil(1 + (breedte_pv + (tussenklem * 10) - 1140) / 940)
    return result


@app.callback(
    Output(component_id='dakgoten', component_property='children'),
    Input(component_id='aantal_rijen_rollen', component_property='children'),
)
def update_output_div(aantal_rijen_rollen):
    result = aantal_rijen_rollen * 2
    return result


@app.callback(
    Output(component_id='schuimstrook_driehoek_profiel', component_property='children'),
    Input(component_id='tussenklem', component_property='children'),
    Input(component_id='lengte_1_rol', component_property='children'),
    Input(component_id='breedte_pv', component_property='children'),
    Input(component_id='eindklem', component_property='children'),
)
def update_output_div(tussenklem, lengte_1_rol, breedte_pv, eindklem):
    result = math.ceil((((breedte_pv + (tussenklem * 10)) * 2) + lengte_1_rol + (eindklem * 10)) / 1280)
    return result


@app.callback(
    Output(component_id='railverbinder', component_property='children'),
    Input(component_id='aantal_rails_van_3_meter_per_rij', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
)
def update_output_div(aantal_rails_van_3_meter_per_rij, aantal_rijen_rails):
    if aantal_rails_van_3_meter_per_rij == 1:
        result = ((aantal_rails_van_3_meter_per_rij - 1) * aantal_rijen_rails)
    else:
        result = ((aantal_rails_van_3_meter_per_rij - 1) * aantal_rijen_rails) + 2
    return result


@app.callback(
    Output(component_id='aantal_ankers_op_1_rail', component_property='children'),
    Input(component_id='lengte_rail', component_property='children'),
    Input(component_id='tussenklem', component_property='children'),
    Input(component_id='anker_plaatsen_om_de', component_property='children')
)
def update_output_div(lengte_rail, tussenklem, anker_plaatsen_om_de):
    result = math.ceil(1 + ((lengte_rail - (20 * tussenklem)) / anker_plaatsen_om_de))
    return result


@app.callback(
    Output(component_id='ankers', component_property='children'),
    Input(component_id='aantal_ankers_op_1_rail', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
)
def update_output_div(aantal_ankers_op_1_rail, aantal_rijen_rails):
    result = math.ceil(aantal_ankers_op_1_rail * aantal_rijen_rails)
    return result


@app.callback(
    Output(component_id='schroeven_voor_ankers', component_property='children'),
    Input(component_id='ankers', component_property='children'),
)
def update_output_div(ankers):
    result = math.ceil(ankers * 3.25)
    return result


@app.callback(
    Output(component_id='beugels', component_property='children'),
    Input(component_id='aantal_ankers_op_1_rail', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
)
def update_output_div(aantal_ankers_op_1_rail, aantal_rijen_rails):
    result = math.ceil(aantal_ankers_op_1_rail * aantal_rijen_rails)
    return result


@app.callback(
    Output(component_id='schroeven_voor_beugels', component_property='children'),
    Input(component_id='aantal_ankers_op_1_rail', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
)
def update_output_div(aantal_ankers_op_1_rail, aantal_rijen_rails):
    return math.ceil(aantal_ankers_op_1_rail * aantal_rijen_rails)


@app.callback(
    Output(component_id='eindklemmen', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
)
def update_output_div(aantal_rijen_rails):
    return aantal_rijen_rails * 2


@app.callback(
    Output(component_id='middenklemmen', component_property='children'),
    Input(component_id='aantal_rijen_rails', component_property='children'),
    Input(component_id='rijen', component_property='value'),
    Input(component_id='kolommen', component_property='value'),
    Input(component_id='indeling', component_property='value')
)
def update_output_div(aantal_rijen_rails, rijen, kolommen, indeling):
    if indeling == 'LND':
        result = aantal_rijen_rails * (rijen - 1)
    else:
        result = aantal_rijen_rails * (kolommen - 1)
    return result


@app.callback(
    Output(component_id='haak', component_property='children'),
    Input(component_id='indeling', component_property='value'),
    Input(component_id='kolommen', component_property='value'),
    Input(component_id='rijen', component_property='value'),
)
def update_output_div(indeling, kolommen, rijen):
    if indeling == 'LND':
        result = (kolommen * 2 * (rijen + 1))
    else:
        result = (rijen * 2 * (kolommen + 1))
    return result


@app.callback(
    Output(component_id='schroeven_voor_hoek', component_property='children'),
    Input(component_id='haak', component_property='children')
)
def update_output_div(haak):
    result = math.ceil(haak * 2.25)
    return result


@app.callback(
    Output(component_id='totaal_aantal_rails_van_3m', component_property='children'),
    Input(component_id='totale_lengte_rails', component_property='children'),
    Input(component_id='raillengte', component_property='children')
)
def update_output_div(totale_lengte_rails, raillengte):
    result = math.ceil(totale_lengte_rails / raillengte) + 1
    return result


@app.callback(
    Output('table', 'children'),
    Output('data', 'children'),
    Output('total_price', 'children'),
    Input('ankers', 'children'),
    Input('totaal_aantal_rails_van_3m', 'children'),
    Input('dakgoten', 'children'),
    Input('schuimstrook_driehoek_profiel', 'children'),
    Input('railverbinder', 'children'),
    Input('schroeven_voor_beugels', 'children'),
    Input('schroeven_voor_ankers', 'children'),
    Input('kleurFrame', 'value'),
    Input('eindklemmen', 'children'),
    Input('middenklemmen', 'children'),
    Input('daksysteem', 'value'),
    Input('aantal_rijen_rollen', 'children')
)
def update_datatable(ankers, totaal_aantal_rails_van_3m, dakgoten,
                     schuimstrook_driehoek_profiel, railverbinder, schroeven_voor_beugels,
                     schroeven_voor_ankers, kleurFrame, eindklemmen, middenklemmen, daksysteem, aantal_rijen_rollen):

    if daksysteem == 'Indak':
        df.loc[df['id'] == "0770001", ['count']] = math.ceil(aantal_rijen_rollen)
        df.loc[df['id'] == "0820239", ['count']] = math.ceil(aantal_rijen_rollen / 100) * 100
    df.loc[df['id'] == "0770003", ['count']] = ankers
    df.loc[df['id'] == "0770212", ['count']] = totaal_aantal_rails_van_3m
    df.loc[df['id'] == "0770037", ['count']] = dakgoten
    df.loc[df['id'] == "0340139", ['count']] = schuimstrook_driehoek_profiel
    df.loc[df['id'] == "0703967", ['count']] = railverbinder
    df.loc[df['id'] == "0770500", ['count']] = math.ceil(schroeven_voor_beugels / 4)
    df.loc[df['id'] == "0770501", ['count']] = math.ceil(schroeven_voor_ankers / 30)
    if kleurFrame == 'ALU':
        df.loc[df['id'] == "0770211", ['count']] = eindklemmen + middenklemmen
    if kleurFrame == 'ALU Zwart':
        df.loc[df['id'] == "0770210", ['count']] = eindklemmen + middenklemmen

    df_result = df.loc[df['count'] > 0].copy()
    df_result['total_price'] = df_result['price'] * df_result['count']
    total_price = "Totale prijs:  {}".format(df_result['total_price'].sum())
    df_result = df_result.astype({'count': 'str', 'total_price': 'str', 'price': 'str'})
    df_result.columns = ['Artikelnummer', 'Omschrijving', 'Bruto', 'Aantal', 'Totaal']
    data = df_result.to_dict('records')
    columns = [{"name": i, "id": i, } for i in df_result.columns]
    return dash_table.DataTable(data=data, columns=columns), json.dumps(data), total_price


@app.callback(
    Output('square', 'children'),
    Input('rijen', 'value'),
    Input('kolommen', 'value')
)
def update_square(rijen=3, kolommen=4):
    encoded_image = base64.b64encode(open(paneel_filename, 'rb').read())

    img_list = []
    for r in range(rijen):
        rows = []
        for k in range(kolommen):
            rows.append(html.Img(src='data:image/png;base64,{}'.format(encoded_image.decode())))
        img_list.append(html.Div(rows, className="row"))

    return img_list
    # return {'width': '40vw', 'height': '10vw'}


@app.callback(
    Output('download-link-docx', 'style'),
    Output('download-link-pdf', 'style'),
    Input('create_advice', 'n_clicks'),
    [
        State('referentie_nr', 'value'),
        State('relatie', 'value'),
        State('data', 'children'),
        State('rijen', 'value'),
        State('kolommen', 'value')
    ]
)
def create_advice(n_clicks, referentie_nr, relatie, json_data, rijen, kolommen):

    docx_button = {'display': 'none'}
    pdf_button = {'display': 'none'}
    if n_clicks > 0:
        # add data to temp advice
        document = MailMerge(template_filename)
        # print(document.get_merge_fields())
        document.merge(
            Relatie=relatie, Referentienummer=referentie_nr
        )
        if json_data is not None:
            data = json.loads(json_data)
            document.merge_rows('Aantal', data)
        document.write(temp_advice_filename)

        # create image for advice
        panel_image = Image.open(paneel_filename)
        new_im = Image.new('RGB', (panel_image.size[0] * kolommen, panel_image.size[1] * rijen))

        for r in range(rijen):
            for c in range(kolommen):
                new_im.paste(panel_image, (c * panel_image.size[0], r * panel_image.size[1]))

        new_im.save(image_filename)

        # add image to advice
        doc = Document(temp_advice_filename)
        tables = doc.tables
        p = tables[0].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        rescale_factor = 1
        r.add_picture(image_filename, width=Inches(kolommen*rescale_factor), height=Inches(rijen*rescale_factor))
        doc.save(advice_filename_docx)
        docx_button = {'display': ''}

        if sys.platform in ('linux'):
            args = ['loffice', '--headless', '--convert-to', 'pdf', '--outdir', './downloads', advice_filename_docx]
            subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            pdf_button = {'display': ''}

    return docx_button, pdf_button


if __name__ == '__main__':
    app.run_server(host='0.0.0.0', port=5050, debug=True)
